import os
import sys
import time
import win32com.client
from pypdf import PdfWriter, PdfReader

# Directory containing the docx files
FOLDER = r"f:\Testing_Tutorial\python_oriented_interview\Printout\selenium"
TEMP_DIR = os.path.join(FOLDER, "temp_pdf_conversion")
OUTPUT_PDF = os.path.join(FOLDER, "Selenium_Complete_Master_Guide.pdf")

# Define logical ordering of the documents with human-readable bookmark titles
DOCUMENTS_IN_ORDER = [
    ("Selenium_WebDriver_Hierarchy.docx", "1. Selenium WebDriver Hierarchy & Architecture"),
    ("WebDriver_RemoteWebDriver.docx", "2. RemoteWebDriver & Driver Core"),
    ("WebDriver_Methods.docx", "3. Core WebDriver Methods"),
    ("WebDriver_Properties.docx", "4. Core WebDriver Properties"),
    ("selenium_method.docx", "5. Essential Selenium Methods"),
    ("selenium_properties.docx", "6. Essential Selenium Properties"),
    ("Selenium_comand.docx", "7. Common Selenium Commands"),
    ("Options_class.docx", "8. Options Class & Browser Capabilities"),
    ("Selenium_waits.docx", "9. Synchronization & Selenium Waits"),
    ("Expected_Conditions.docx", "10. Expected Conditions Reference"),
    ("ActionChain_class.docx", "11. ActionChains Class (Advanced User Interactions)"),
    ("Selenium_Special_Keys.docx", "12. Selenium Special Keys & Shortcuts"),
    ("drop_down_selection.docx", "13. Dropdown & Select Element Handling"),
    ("Alerts_popup_handle.docx", "14. Handling Alerts & Popups"),
    ("frames_iframes.docx", "15. Handling Frames & Iframes"),
    ("handle_window_tab.docx", "16. Managing Windows & Browser Tabs"),
    ("Handling_file_uploads.docx", "17. Handling File Uploads"),
    ("python_javascript_executor.docx", "18. JavaScript Executor in Python"),
    ("JAVASCRIPTEXECUTOR IN SELENIUM WITH JAVA.docx", "19. JavaScript Executor in Java"),
    ("screenshots_python_selenium.docx", "20. Capturing Screenshots in Selenium"),
    ("validate_broken_links.docx", "21. Validating Broken Links"),
    ("Capturing_network_logs_Selenium_Python.docx", "22. Capturing Network Logs in Selenium"),
    ("exceptions_selenium_pytest.docx", "23. Selenium & Pytest Exception Handling"),
]

def create_cover_page_docx(path):
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = docx.Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(72)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("SELENIUM AUTOMATION & TESTING\nCOMPREHENSIVE MASTER GUIDE")
    run_title.bold = True
    run_title.font.size = Pt(26)
    run_title.font.name = "Calibri"
    run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)  # Deep Blue

    # Divider line
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.paragraph_format.space_after = Pt(24)
    r_div = p_div.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    r_div.font.color.rgb = RGBColor(0x2F, 0x55, 0x97)
    r_div.font.size = Pt(14)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(48)
    run_sub = p_sub.add_run("Complete Collection of 23 Core Topics: Architecture, WebDriver API, Waits,\nActionChains, Window & Frame Handling, JS Execution, Network Logs & Exceptions")
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.name = "Calibri"
    run_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # Details Box / Summary Table
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(36)
    p_meta.paragraph_format.space_after = Pt(12)
    r_meta = p_meta.add_run("📘 DOCUMENT SUMMARY")
    r_meta.bold = True
    r_meta.font.size = Pt(14)
    r_meta.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    data = [
        ("Total Modules Merged:", f"{len(DOCUMENTS_IN_ORDER)} Docx Documents"),
        ("Primary Domain:", "Selenium WebDriver (Python & Java)"),
        ("Features Included:", "Full Code Snippets, Hierarchy, Waits & Exceptions"),
        ("Generated Date:", time.strftime("%B %d, %Y")),
    ]

    for row_idx, (k, v) in enumerate(data):
        row = table.rows[row_idx]
        cell_k, cell_v = row.cells[0], row.cells[1]
        
        pk = cell_k.paragraphs[0]
        pk.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rk = pk.add_run(k)
        rk.bold = True
        rk.font.size = Pt(11)
        rk.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        pv = cell_v.paragraphs[0]
        pv.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rv = pv.add_run(v)
        rv.font.size = Pt(11)
        rv.font.color.rgb = RGBColor(0x2F, 0x55, 0x97)

    doc.save(path)
    print("Created cover page docx successfully.")

def convert_docx_to_pdf(word_app, docx_path, pdf_path):
    docx_abs = os.path.abspath(docx_path)
    pdf_abs = os.path.abspath(pdf_path)
    
    # 17 = wdFormatPDF
    doc = word_app.Documents.Open(docx_abs, ReadOnly=True)
    doc.SaveAs(pdf_abs, FileFormat=17)
    doc.Close()

def main():
    if not os.path.exists(TEMP_DIR):
        os.makedirs(TEMP_DIR)
        
    cover_docx = os.path.join(TEMP_DIR, "00_Cover_Page.docx")
    cover_pdf = os.path.join(TEMP_DIR, "00_Cover_Page.pdf")
    
    create_cover_page_docx(cover_docx)
    
    print("\nStarting Word COM Application...")
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    
    converted_pdfs = []
    
    try:
        # 1. Convert Cover Page
        print("Converting Cover Page to PDF...")
        convert_docx_to_pdf(word, cover_docx, cover_pdf)
        converted_pdfs.append((cover_pdf, "Cover Page"))
        
        # 2. Convert all docx files to PDF
        for idx, (filename, title) in enumerate(DOCUMENTS_IN_ORDER, 1):
            docx_path = os.path.join(FOLDER, filename)
            if not os.path.exists(docx_path):
                print(f"WARNING: File {filename} not found! Skipping...")
                continue
                
            pdf_name = f"{idx:02d}_{os.path.splitext(filename)[0]}.pdf"
            pdf_path = os.path.join(TEMP_DIR, pdf_name)
            
            print(f"[{idx}/{len(DOCUMENTS_IN_ORDER)}] Converting {filename} -> PDF...")
            convert_docx_to_pdf(word, docx_path, pdf_path)
            converted_pdfs.append((pdf_path, title))
            
    finally:
        word.Quit()
        print("Closed Word COM Application.")

    # 3. Merge all PDFs using PyPDF with Bookmarks
    print("\nMerging all PDFs into single master PDF with Bookmarks...")
    writer = PdfWriter()
    
    current_page = 0
    for pdf_path, bookmark_title in converted_pdfs:
        reader = PdfReader(pdf_path)
        num_pages = len(reader.pages)
        
        # Append pages
        for page in reader.pages:
            writer.add_page(page)
            
        # Add Bookmark
        writer.add_outline_item(bookmark_title, current_page)
        current_page += num_pages
        print(f"Added bookmark '{bookmark_title}' at page {current_page - num_pages + 1} (total pages so far: {current_page})")

    # Save output PDF
    with open(OUTPUT_PDF, "wb") as f_out:
        writer.write(f_out)

    print(f"\nSUCCESS! Master PDF created at:\n{OUTPUT_PDF}")
    print(f"Total Pages in merged PDF: {current_page}")

if __name__ == "__main__":
    main()
